"""
utils.py — Text extraction and file hashing utilities.
Supports PDF, DOCX, and TXT files.
"""

import hashlib
import os
import re


# ─────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────

def extract_text(filepath: str) -> str:
    """
    Extract and clean text from PDF, DOCX, or TXT files.

    Cleaning steps applied to all formats:
      - Remove empty / whitespace-only lines
      - Normalize internal whitespace (collapse tabs, multiple spaces)
      - Strip leading/trailing whitespace per line

    Args:
        filepath: Absolute or relative path to the document.

    Returns:
        Cleaned text as a single string.

    Raises:
        ValueError: If the file extension is not supported.
        FileNotFoundError: If the file does not exist.
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        raw_text = _extract_pdf(filepath)
    elif ext == ".docx":
        raw_text = _extract_docx(filepath)
    elif ext == ".txt":
        raw_text = _extract_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: .pdf, .docx, .txt")

    return _clean_text(raw_text)


def _extract_pdf(filepath: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF extraction. Run: pip install pdfplumber")

    pages = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                content = page.extract_text()
                if content:
                    pages.append(content)
            except Exception as e:
                print(f"  [Warning] Could not extract page {i+1} of '{filepath}': {e}")
    return "\n".join(pages)


def _extract_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX extraction. Run: pip install python-docx")

    doc = Document(filepath)
    paragraphs = [para.text for para in doc.paragraphs]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def _extract_txt(filepath: str) -> str:
    """Extract text from a plain text file with encoding fallback."""
    encodings = ["utf-8", "latin-1", "cp1252", "ascii"]
    for enc in encodings:
        try:
            with open(filepath, "r", encoding=enc, errors="strict") as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    # Final fallback: replace undecodable characters
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _clean_text(text: str) -> str:
    """
    Clean raw extracted text:
      1. Normalize line endings
      2. Strip each line and collapse internal whitespace
      3. Remove empty lines
    """
    # Normalize Windows/Mac line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    cleaned_lines = []
    for line in text.splitlines():
        # Collapse tabs and multiple spaces to a single space
        line = re.sub(r"[ \t]+", " ", line).strip()
        if line:  # Drop empty lines
            cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


# ─────────────────────────────────────────────
# FILE METADATA & HASHING
# ─────────────────────────────────────────────

def file_hash(filepath: str) -> str:
    """
    Compute the SHA-256 hash of a file's binary content.
    Used for exact duplicate detection.

    Args:
        filepath: Path to the file.

    Returns:
        Hex-encoded SHA-256 digest string.
    """
    sha256 = hashlib.sha256()
    with open(filepath, "rb") as f:
        # Read in chunks to handle large files efficiently
        for chunk in iter(lambda: f.read(65536), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def file_metadata(filepath: str) -> dict:
    """
    Collect file metadata for storage alongside the document.

    Returns a dict with:
      - file_path   : absolute path
      - file_size   : size in bytes
      - modified_time: last-modified timestamp (Unix epoch float)
      - extension   : lowercase file extension
    """
    stat = os.stat(filepath)
    return {
        "file_path": os.path.abspath(filepath),
        "file_size": stat.st_size,
        "modified_time": stat.st_mtime,
        "extension": os.path.splitext(filepath)[1].lower(),
    }